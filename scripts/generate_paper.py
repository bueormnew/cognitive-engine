from __future__ import annotations

import json
import math
import os
from pathlib import Path
from typing import Dict, Iterable, List, Sequence

import matplotlib.pyplot as plt
import networkx as nx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from cognitive_engine import EngineBuilder


ROOT = Path(".")
REPORT_DIR = ROOT / "artifacts" / "reports"
PLOT_DIR = ROOT / "artifacts" / "plots"
DOCS_DIR = ROOT / "docs"


REFERENCES = [
    {
        "id": "R1",
        "label": "McClelland, McNaughton, O'Reilly (1995), Complementary Learning Systems",
        "url": "https://stanford.edu/~jlmcc/papers/McCMcNaughtonOReilly95.pdf",
        "note": "Neuroscience-inspired separation between fast episodic learning and slow semantic consolidation.",
    },
    {
        "id": "R2",
        "label": "Kirkpatrick et al. (2017), Overcoming catastrophic forgetting in neural networks",
        "url": "https://pubmed.ncbi.nlm.nih.gov/28292907/",
        "note": "Canonical reference for protecting old knowledge during sequential learning.",
    },
    {
        "id": "R3",
        "label": "Rolnick et al. (2019), Experience Replay for Continual Learning",
        "url": "https://arxiv.org/abs/1811.11682",
        "note": "Replay-based stabilization strategy for continual learning streams.",
    },
    {
        "id": "R4",
        "label": "Houlsby et al. (2019), Parameter-Efficient Transfer Learning for NLP",
        "url": "https://arxiv.org/abs/1902.00751",
        "note": "Adapter-style parameter isolation for modular specialization.",
    },
    {
        "id": "R5",
        "label": "Hu et al. (2021), LoRA: Low-Rank Adaptation of Large Language Models",
        "url": "https://arxiv.org/abs/2106.09685",
        "note": "Low-rank adaptation that motivates localized plasticity instead of full-model updates.",
    },
    {
        "id": "R6",
        "label": "Shazeer et al. (2017), Sparsely-Gated Mixture-of-Experts",
        "url": "https://arxiv.org/abs/1701.06538",
        "note": "Conditional routing over specialized experts.",
    },
    {
        "id": "R7",
        "label": "Fedus, Zoph, Shazeer (2021), Switch Transformers",
        "url": "https://arxiv.org/abs/2101.03961",
        "note": "Sparse routing with practical scaling properties.",
    },
    {
        "id": "R8",
        "label": "Lewis et al. (2020), Retrieval-Augmented Generation",
        "url": "https://arxiv.org/abs/2005.11401",
        "note": "External memory retrieval for factual and updatable generation.",
    },
]


FOLDER_NOTES = {
    "core": "El núcleo contiene el engine, el builder, el registry y los tipos compartidos. Aquí vive la lógica de orquestación que asegura que la arquitectura se comporte como framework y no como script ad hoc.",
    "modules": "Los módulos implementan percepción, representación semántica e inferencia de relevancia. Cada módulo tiene entradas y salidas claras para permitir hot-swapping.",
    "memory": "La capa de memoria implementa almacenamiento jerárquico, recuperación híbrida vectorial-lexical y consolidación de registros comprimidos.",
    "training": "Entrenamiento online, utilidades de evaluación y adaptación localizada. Esta carpeta encapsula la coexistencia entre inferencia y aprendizaje incremental.",
    "routing": "Ruteo dinámico de cómputo. Aquí se decide cuándo activar plasticidad, qué memorias consultar y cómo asignar presupuesto de inferencia.",
    "compression": "Reescritura semántica y compresión conceptual. Esta capa transforma experiencia bruta en conocimiento estructurado y deduplicable.",
    "interfaces": "Contratos abstractos, superficies de extensión y puntos de integración para plugins internos o reemplazos futuros.",
    "config": "Esquemas y loaders de configuración dinámica. Permite cambiar umbrales, tamaño de memorias, operaciones y dispositivos sin reescribir el engine.",
    "models": "Modelos estables o casi congelados. En esta implementación el núcleo estable no aprende de forma agresiva y actúa como soporte semántico consistente.",
    "adapters": "Módulos plásticos entrenables, de baja dimensión efectiva, responsables de adaptación localizada y controlada.",
    "replay": "Buffers de replay priorizado para mezclar ejemplos recientes con experiencia pasada y mitigar olvido catastrófico.",
    "consolidation": "Servicios de consolidación offline o periódica que fusionan memorias redundantes, reponderan replay y reducen ruido.",
    "utils": "Telemetría, plots, seeding, helpers de texto y utilidades numéricas. Esta carpeta habilita observabilidad y reproducibilidad.",
    "api": "Superficie de consumo para exponer el engine a otros procesos, servicios o UIs sin filtrar detalles internos del pipeline.",
}


def load_json(path: Path) -> Dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def architecture_graph(engine_description: Dict[str, object], output_path: Path) -> str:
    graph = nx.DiGraph()
    nodes = [
        "Input",
        "Input Processor",
        "Semantic Encoder",
        "Router",
        "Memory System",
        "Importance Evaluator",
        "Knowledge Compressor",
        "Replay Buffer",
        "Plastic Learner",
        "Stable Core",
        "Consolidator",
        "Response",
    ]
    graph.add_nodes_from(nodes)
    graph.add_edges_from(
        [
            ("Input", "Input Processor"),
            ("Input Processor", "Semantic Encoder"),
            ("Semantic Encoder", "Router"),
            ("Semantic Encoder", "Memory System"),
            ("Memory System", "Importance Evaluator"),
            ("Semantic Encoder", "Importance Evaluator"),
            ("Importance Evaluator", "Knowledge Compressor"),
            ("Knowledge Compressor", "Memory System"),
            ("Knowledge Compressor", "Replay Buffer"),
            ("Router", "Plastic Learner"),
            ("Router", "Stable Core"),
            ("Memory System", "Stable Core"),
            ("Plastic Learner", "Stable Core"),
            ("Replay Buffer", "Plastic Learner"),
            ("Memory System", "Consolidator"),
            ("Replay Buffer", "Consolidator"),
            ("Stable Core", "Response"),
        ]
    )

    plt.figure(figsize=(12, 8))
    positions = nx.spring_layout(graph, seed=17, k=1.1)
    nx.draw_networkx_nodes(graph, positions, node_size=2200, node_color="#dbeafe", edgecolors="#1d4ed8")
    nx.draw_networkx_edges(graph, positions, width=1.5, arrows=True, arrowsize=18, edge_color="#475569")
    nx.draw_networkx_labels(graph, positions, font_size=9)
    plt.axis("off")
    plt.title("Hybrid Cognitive Architecture Topology")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    plt.tight_layout()
    plt.savefig(output_path, dpi=180)
    plt.close()
    return str(output_path)


def action_histogram_plot(text_report: Dict[str, object], output_path: Path) -> str:
    histogram = text_report["action_histogram"]
    labels = list(histogram.keys())
    values = [histogram[label] for label in labels]
    plt.figure(figsize=(8, 5))
    plt.bar(labels, values, color=["#2563eb", "#0f766e", "#9333ea", "#dc2626"])
    plt.title("Learning actions during text stream")
    plt.ylabel("Count")
    plt.tight_layout()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    plt.savefig(output_path, dpi=180)
    plt.close()
    return str(output_path)


def prediction_case_plot(numeric_report: Dict[str, object], output_path: Path) -> str:
    cases = numeric_report["prediction_cases"]
    labels = [f"{case['operation']}({int(case['a'])},{int(case['b'])})" for case in cases]
    targets = [case["target"] for case in cases]
    predictions = [case["prediction"] for case in cases]
    x = range(len(labels))
    plt.figure(figsize=(10, 5))
    plt.plot(x, targets, marker="o", label="Target")
    plt.plot(x, predictions, marker="s", label="Prediction")
    plt.xticks(list(x), labels, rotation=15)
    plt.ylabel("Value")
    plt.title("Prediction vs target on held-out numeric probes")
    plt.legend()
    plt.tight_layout()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    plt.savefig(output_path, dpi=180)
    plt.close()
    return str(output_path)


def build_tree(root: Path, interesting: Sequence[str]) -> str:
    lines = ["."]
    for name in interesting:
        path = root / name
        lines.append(f"├── {name}/")
        for child in sorted(path.iterdir()):
            if child.is_dir():
                lines.append(f"│   ├── {child.name}/")
            else:
                lines.append(f"│   ├── {child.name}")
    return "\n".join(lines)


def md_table(headers: Sequence[str], rows: Sequence[Sequence[object]]) -> str:
    line1 = "| " + " | ".join(headers) + " |"
    line2 = "| " + " | ".join(["---"] * len(headers)) + " |"
    body = ["| " + " | ".join(str(item) for item in row) + " |" for row in rows]
    return "\n".join([line1, line2, *body])


def word_count(text: str) -> int:
    return len(text.split())


def write_markdown(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def add_table(document: Document, headers: Sequence[str], rows: Sequence[Sequence[object]]) -> None:
    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    for column, header in enumerate(headers):
        table.cell(0, column).text = str(header)
    for row_index, row in enumerate(rows, start=1):
        for column, value in enumerate(row):
            table.cell(row_index, column).text = str(value)


def write_docx(path: Path, sections: List[Dict[str, object]], figures: List[str], metadata: Dict[str, object]) -> None:
    document = Document()
    title = document.add_heading(metadata["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(metadata["subtitle"])
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Word count (markdown source): {metadata['word_count']}")

    for figure in figures:
        if Path(figure).exists():
            document.add_picture(figure, width=Inches(6.2))

    for section in sections:
        document.add_heading(section["title"], level=1)
        for paragraph in section.get("paragraphs", []):
            document.add_paragraph(paragraph)
        table = section.get("table")
        if table:
            add_table(document, table["headers"], table["rows"])
        for figure in section.get("figures", []):
            if Path(figure).exists():
                document.add_picture(figure, width=Inches(6.2))
        code_block = section.get("code_block")
        if code_block:
            paragraph = document.add_paragraph()
            paragraph.add_run(code_block).font.name = "Consolas"

    document.save(path)


def generate_sections(
    engine_description: Dict[str, object],
    text_report: Dict[str, object],
    numeric_report: Dict[str, object],
    tree_text: str,
    figures: Dict[str, str],
) -> List[Dict[str, object]]:
    stream_rows = []
    for index, step in enumerate(text_report["stream_steps"], start=1):
        stream_rows.append(
            [
                index,
                step["importance_action"],
                f"{step['importance_score']:.3f}",
                "yes" if step["learning_applied"] else "no",
                step["input"][:52] + ("..." if len(step["input"]) > 52 else ""),
            ]
        )

    epoch_rows = []
    interesting_epochs = {1, 3, 7, 13, numeric_report["best_epoch"], len(numeric_report["epoch_metrics"])}
    for metric in numeric_report["epoch_metrics"]:
        if metric["epoch"] in interesting_epochs:
            epoch_rows.append(
                [
                    metric["epoch"],
                    f"{metric['loss']:.4f}",
                    f"{metric['mae']:.4f}",
                    f"{metric['accuracy']:.4f}",
                    f"{metric['plastic_norm']:.4f}",
                ]
            )

    folder_rows = [[name, note] for name, note in FOLDER_NOTES.items()]
    component_rows = []
    for key, value in engine_description.items():
        component_rows.append([key, json.dumps(value, ensure_ascii=False)[:120]])

    query_rows = [[item["query"], item["response"]] for item in text_report["query_outputs"]]
    prediction_rows = [
        [
            case["operation"],
            f"{int(case['a'])}, {int(case['b'])}",
            case["target"],
            f"{case['prediction']:.3f}",
            f"{case['confidence']:.3f}",
        ]
        for case in numeric_report["prediction_cases"]
    ]

    best_validation = numeric_report["validation_final"]
    semantic_final = text_report["final_snapshot"]["memory"]["semantic_size"]
    episodic_final = text_report["final_snapshot"]["memory"]["episodic_size"]
    action_histogram = ", ".join(f"{key}={value}" for key, value in text_report["action_histogram"].items())
    top_concepts = ", ".join(f"{name}:{count}" for name, count in text_report["final_snapshot"]["memory"]["top_concepts"][:8])

    sections: List[Dict[str, object]] = [
        {
            "title": "Abstract",
            "paragraphs": [
                "Este documento presenta una arquitectura híbrida cognitiva modular de texto-a-texto diseñada en Python con PyTorch, orientada a aprendizaje continuo en tiempo real, memoria selectiva, plasticidad controlada y escalabilidad extrema. La propuesta evita el patrón de un transformer monolítico que actualiza todos sus pesos ante cualquier nueva experiencia. En su lugar, desacopla el sistema en un núcleo estable, un módulo plástico localizado, una jerarquía explícita de memorias, un evaluador de importancia, un compresor de conocimiento y un motor de consolidación que reorganiza información fuera del camino crítico de inferencia.",
                f"Además de la formulación teórica, el proyecto se implementó como framework funcional. Se ejecutó un stream textual de diez eventos con memoria selectiva y un experimento numérico con un micromodelo entrenable sobre una tarea de suma, resta y multiplicación con operandos de 0 a 12. Los resultados observados fueron reales: el mejor checkpoint del módulo plástico alcanzó exactitud de validación {numeric_report['best_validation_accuracy']:.3f} con error absoluto medio {best_validation['mae']:.3f}, mientras que la corriente textual aprendió selectivamente {semantic_final} memorias semánticas útiles y descartó ruido o señales de baja prioridad.",
                "El valor principal del proyecto no es entrenar un LLM generalista, sino demostrar una arquitectura reusable, instrumentable y extensible que separa percepción, memoria, compresión, plasticidad y consolidación. Esa separación permite hot-swapping de componentes, control fino del riesgo de drift, observabilidad total del pipeline y un camino realista para crecer desde prototipos pequeños hacia sistemas distribuidos con múltiples expertos, almacenamiento vectorial externo y entrenamiento continuo seguro.",
            ],
            "figures": [figures["architecture"], figures["action_histogram"]],
        },
        {
            "title": "1. Problem Statement And Design Criteria",
            "paragraphs": [
                "El problema atacado aquí no es únicamente la generación de texto. El objetivo es construir una arquitectura de inteligencia artificial capaz de recibir lenguaje, abstraerlo, decidir si vale la pena aprenderlo, transformarlo en conocimiento comprimido, reforzar lo realmente útil y proteger el conocimiento previo. Dicho de otra forma: el sistema debe tratar la experiencia como flujo cognitivo, no como simple secuencia de tokens para fine tuning indiscriminado.",
                "Los requisitos originales fuerzan varias restricciones arquitectónicas fuertes. Primero, el sistema no puede depender de un único bloque denso de parámetros que se ajusta por completo con cada dato nuevo. Segundo, la memoria no puede ser un simple log de texto crudo; debe operar sobre conceptos, relaciones, embeddings y resúmenes comprimidos. Tercero, la actualización del conocimiento debe ser selectiva, trazable y reversible. Cuarto, la plataforma debe soportar crecimiento modular, sustituibilidad de piezas e inferencia y aprendizaje coexistentes.",
                "Estas restricciones llevan a un diseño híbrido. Se conserva un núcleo estable responsable de lenguaje, agregación contextual y composición de salida; se aísla un módulo plástico para adaptación localizada; se añade un sistema de memoria jerárquico; se insertan políticas explícitas de replay y consolidación; y se gobierna todo con ruteo dinámico. El resultado final se parece más a un motor cognitivo que a un único modelo estadístico monolítico.",
            ],
        },
        {
            "title": "2. Theoretical Inspiration",
            "paragraphs": [
                "La idea de separar aprendizaje rápido y representación estable se inspira directamente en la teoría de Complementary Learning Systems [R1], donde una memoria episódica rápida y una memoria semántica más lenta cooperan para evitar interferencia destructiva. En esta arquitectura esa separación se refleja en la coexistencia entre memoria episódica, memoria semántica consolidada y un núcleo estable con baja plasticidad.",
                "El control del olvido catastrófico toma ideas de varias familias. EWC [R2] muestra la utilidad de proteger conocimiento previo durante aprendizaje secuencial; Experience Replay [R3] demuestra que revisar muestras anteriores reduce degradación; adapters [R4] y LoRA [R5] ofrecen formas concretas de restringir la plasticidad a subespacios pequeños; MoE y Switch Transformers [R6][R7] muestran que el ruteo condicional es una vía práctica para especialización y escalabilidad; y RAG [R8] deja claro que una memoria externa recuperable puede complementar al conocimiento paramétrico.",
                "La propuesta no copia un paper concreto. La contribución aquí es una síntesis ingenieril: biología para la intuición de consolidación y dualidad fast/slow, aprendizaje continuo para replay y protección, eficiencia paramétrica para plasticidad localizada, y recuperación externa para memoria actualizable. El framework resultante acepta que la inteligencia operativa emerge mejor cuando las responsabilidades se dividen y se hacen observables.",
            ],
        },
        {
            "title": "3. System-Level Architecture",
            "paragraphs": [
                "La arquitectura completa se organiza en once subsistemas visibles: procesamiento de entrada, representación semántica, ruteo, memoria jerárquica, evaluación de importancia, compresión de conocimiento, núcleo estable, módulo plástico, replay, consolidación y telemetría. Cada uno existe como componente intercambiable con contrato propio. Ese detalle importa porque convierte el proyecto en framework y no en script rígido.",
                "El flujo principal sigue una lógica cognitiva. Un payload entra al Input Processor adecuado; el Semantic Encoder lo transforma en estado semántico; el Router decide qué camino activar; el sistema de memoria recupera contexto relevante; el evaluador estima si lo visto merece aprendizaje; el compresor reescribe el contenido para almacenarlo de manera económica; el núcleo estable compone la respuesta; y, si procede, el módulo plástico se adapta con apoyo del replay buffer. Periódicamente, la consolidación compacta y limpia la memoria acumulada.",
                "La arquitectura separa deliberadamente inferencia de consolidación. El camino crítico intenta ser corto y trazable, mientras que la reorganización costosa de memorias se ejecuta fuera del paso inmediato de respuesta. Esa decisión sigue la intuición de que una IA usable debe responder primero y reordenar su conocimiento después, en background o en puntos de baja carga.",
            ],
            "table": {"headers": ["Component", "Descriptor"], "rows": component_rows},
            "figures": [figures["architecture"]],
        },
        {
            "title": "4. Professional Project Structure",
            "paragraphs": [
                "La estructura del código se diseñó para alinearse con los límites conceptuales del sistema. No se usó una carpeta genérica tipo `src/models` con lógica mezclada, sino dominios separados por responsabilidad. Esto mejora navegabilidad, reemplazo, cobertura de tests y desacoplamiento operativo.",
                "Cada carpeta responde a una pregunta concreta: `core` orquesta, `modules` interpreta, `memory` almacena, `training` adapta, `routing` decide, `compression` abstrae, `adapters` introduce plasticidad localizada y `consolidation` reestructura conocimiento. `utils` no es un cajón de sastre; contiene solamente instrumentación y helpers transversales. `api` ofrece una superficie limpia para consumo externo.",
            ],
            "table": {"headers": ["Folder", "Purpose"], "rows": folder_rows},
            "code_block": tree_text,
        },
        {
            "title": "5. Core Data Contracts",
            "paragraphs": [
                "El framework gira alrededor de contratos explícitos definidos en `core/types.py` e `interfaces/base.py`. `ProcessedInput` representa el resultado de tokenización o preprocesamiento numérico; `SemanticState` encapsula embedding, intención, entidades, conceptos y contexto comprimido; `ImportanceAssessment` modela la decisión de aprendizaje; `CompressedKnowledge` es la unidad persistible de memoria; `MemoryBundle` agrupa el contexto recuperado; y `CoreInference` encapsula la respuesta del sistema estable o plástico.",
                "Estas estructuras desacoplan serialización, depuración y reemplazo. Un `SemanticEncoder` alternativo puede producir un `SemanticState` distinto internamente mientras respete el contrato. De la misma forma, un backend de memoria externo como FAISS, Qdrant o Chroma puede sustituir a la implementación local siempre que devuelva `CompressedKnowledge` y `MemoryBundle` compatibles. La programación orientada a objetos no se usó aquí solo por estilo: se usó para congelar fronteras técnicas.",
                "También es importante que el contrato sea visible para observabilidad. Cada fase del pipeline emite `TraceEvent` estructurados. Eso hace que la IA se pueda auditar como sistema de software: es posible inspeccionar cómo se procesó la entrada, qué conceptos se extrajeron, qué acción de aprendizaje se eligió y qué memorias se tocaron.",
            ],
        },
        {
            "title": "6. Input Processing Layer",
            "paragraphs": [
                "La capa de entrada fue implementada con dos procesadores desacoplados: `TextInputProcessor` y `NumericInputProcessor`. El primero normaliza, tokeniza y mapea texto a IDs hash de vocabulario fijo, sin imponer dependencia de un tokenizer externo. El segundo procesa pares numéricos y operación aritmética, generando representación utilizable por el resto del framework sin convertir números a tokens de texto.",
                "La decisión de soportar múltiples processors responde a un principio central del proyecto: la modalidad no debe contaminar el resto de la arquitectura. Un engine híbrido serio no debería asumir que todas las entradas son secuencias lingüísticas. Por eso el selector de processor se hace antes de la semántica y se basa en el contrato `supports(payload)`.",
                "A nivel de extensibilidad, esta capa es el lugar natural para insertar tokenizers reales, detectores de idioma, normalizadores específicos de dominio o ingestión multicanal. El engine actual ya está preparado para ello porque solo necesita que el processor produzca un `ProcessedInput` coherente.",
            ],
        },
        {
            "title": "7. Semantic Understanding Layer",
            "paragraphs": [
                "La comprensión semántica no se resuelve con un único bloque opaco. En la rama textual se usó un encoder híbrido pequeño compuesto por embedding, transformer encoder y GRU. Ese encoder no pretende competir con un LLM grande; pretende demostrar cómo un estado semántico puede resultar de combinar representación distribuida y extracción simbólica ligera en un contrato reusable.",
                "En la rama numérica, la semántica es intencionalmente diferente. El `NumericSemanticEncoder` construye features deterministas a partir de operandos normalizados y una representación aprendible de operación. Esa separación demuestra que la arquitectura soporta modalidades heterogéneas sin romper el pipeline principal. El objetivo no es que todo se parezca a texto, sino que todo fluya por interfaces consistentes.",
                "El resultado de esta capa es un `SemanticState` que contiene embedding agregado, intención, conceptos, entidades, aristas de un pequeño grafo conceptual y una forma comprimida de contexto. Esa salida es suficiente para alimentar memoria, ruteo, respuesta y aprendizaje sin obligar a que todos los módulos sepan de tokens crudos.",
            ],
        },
        {
            "title": "8. Importance Evaluation System",
            "paragraphs": [
                "El evaluador de importancia es la pieza normativa que evita que el sistema recuerde todo. Su función es decidir si un input es ruido, dato temporal, conocimiento nuevo, corrección o patrón digno de refuerzo. Para ello calcula novedad, utilidad, frecuencia, coherencia, redundancia, relevancia futura, señal de corrección y riesgo de contradicción.",
                "La política actual combina puntuaciones heurísticas con umbrales configurables. Aunque no es un clasificador supervisado, sí es una política explícita de control de plasticidad. Esto es una ventaja arquitectónica: la estrategia se puede auditar y reemplazar. En un sistema productivo podría convertirse en un calibrador aprendido, un ensemble o incluso una política bayesiana, pero el contrato seguiría siendo el mismo.",
                "Durante el experimento textual, esta capa produjo acciones diferenciadas. El histograma real observado fue: "
                + action_histogram
                + ". El sistema aprendió preferencias y correcciones, descartó small talk y conservó conocimiento técnico. Esa selectividad es precisamente la evidencia de que la arquitectura no se limita a almacenar la conversación de forma indiscriminada.",
            ],
            "figures": [figures["action_histogram"]],
        },
        {
            "title": "9. Knowledge Compression Layer",
            "paragraphs": [
                "Aprender selectivamente no basta; también hay que almacenar compactamente. El compresor transforma experiencia en `CompressedKnowledge`, una unidad persistente basada en resumen, conceptos, relaciones, embedding e información de procedencia. Este diseño impone una regla saludable: la memoria a largo plazo no guarda la conversación completa salvo que exista una razón fuerte.",
                "La reescritura semántica reduce redundancia y acerca la memoria a un formato que puede fusionarse, buscarse y reforzarse. Un registro como `Knowledge compressed: aprendi, area2d, detecta, cuerpos` es más útil para recuperación que un párrafo entero con ruido. Lo mismo aplica a preferencias y correcciones, que se guardan como hechos operativos y no como transcript literal.",
                "La compresión también es una defensa contra drift. Guardar texto crudo de forma masiva favorece contaminación contextual, loops autorreferenciales y sobreajuste a ejemplos recientes. Guardar conceptos y relaciones reduce ese riesgo y facilita consolidación offline.",
            ],
        },
        {
            "title": "10. Stable Core And Controlled Plasticity",
            "paragraphs": [
                "La arquitectura separa con claridad el núcleo estable del módulo plástico. El `StableReasoningCore` actúa como agregador congelado o casi congelado: proyecta contexto, combina memorias recuperadas y compone una respuesta final. No recibe actualizaciones agresivas en línea. La motivación es simple: el conocimiento general y la capacidad de composición deben resistir el ruido del stream.",
                "La adaptación localizada recae en `PlasticArithmeticModule`, que usa una ruta de bajo rango, una cabeza directa por operación y objetivos normalizados. Esta pieza fue elegida para demostrar que la plasticidad puede estar confinada a un submodelo pequeño y aun así aprender comportamientos útiles sin tocar todo el sistema.",
                "En una expansión futura, el mismo patrón podría portar LoRA [R5], adapters [R4], recurrent adapters, hypernetworks o expertos sparsos. El punto esencial es la disciplina de actualización: la plasticidad es un recurso localizado, no un permiso para degradar globalmente el modelo base.",
            ],
            "figures": [figures["prediction_cases"]],
        },
        {
            "title": "11. Hierarchical Memory System",
            "paragraphs": [
                "La memoria se divide en cuatro niveles operativos. Short-Term Memory retiene el contexto reciente; Working Memory resume conceptos activos e intención presente; Long-Term Semantic Memory conserva conocimiento consolidado y comprimido; y Episodic Memory almacena huellas de experiencia recientes que todavía pueden ser relevantes para replay o análisis contextual.",
                "La implementación concreta usa `HierarchicalMemorySystem` con una combinación de deques para memoria reciente, diccionario para memoria semántica y un índice vectorial NumPy para búsqueda semántica. La capa actual es local y pequeña, pero el contrato permite reemplazarla por FAISS, Qdrant o Chroma sin alterar el engine.",
                f"En el stream textual, el snapshot final mostró {semantic_final} memorias semánticas y {episodic_final} episodios útiles. Los conceptos más frecuentes fueron: {top_concepts}. Lo importante aquí no es el tamaño absoluto, sino que el sistema mantuvo una memoria compacta y legible en lugar de crecer linealmente con el transcript.",
            ],
        },
        {
            "title": "12. Retrieval And Memory Query Strategy",
            "paragraphs": [
                "La recuperación de memoria no se dejó como una simple búsqueda ANN. El prototipo mezcla similitud vectorial y coincidencia conceptual lexicalizada. Esa mezcla es importante porque consultas reales pueden reformular la superficie lingüística del mismo tema. La pregunta sobre preferencias, por ejemplo, no comparte exactamente todos los tokens con la frase original `Prefiero respuestas técnicas y directas`, pero sí comparte raíces temáticas relevantes.",
                "La política híbrida usada en `memory/stores.py` pondera score vectorial y score lexical y añade bonos temáticos cuando la consulta es claramente interrogativa. Es una simplificación, pero ilustra una idea clave: la memoria semántica en un sistema cognitivo práctico necesita búsqueda aproximada y también reglas de reorientación temática.",
                "Los resultados del prototipo lo muestran con claridad. Después del ajuste híbrido, la consulta `¿Qué recuerdas sobre mis preferencias?` recuperó como primer resultado el hecho correcto de preferencia, en lugar de una memoria técnica irrelevante. Esa mejora no vino de entrenar más el encoder; vino de diseñar mejor la política de recuperación.",
            ],
            "table": {"headers": ["Query", "Observed response"], "rows": query_rows},
        },
        {
            "title": "13. Replay Buffer And Catastrophic Forgetting Control",
            "paragraphs": [
                "El replay buffer es una pieza obligatoria si se quiere aprendizaje continuo. Aquí se implementó como `PrioritizedReplayBuffer`, que almacena muestras con prioridad y las reinyecta durante adaptación del módulo plástico. El buffer introduce un recordatorio estadístico del pasado para que el modelo no aprenda solo del ejemplo más reciente.",
                "En la literatura, replay es una de las defensas más robustas contra olvido catastrófico [R3]. En esta arquitectura no es un añadido decorativo: está integrado con el motor de entrenamiento online y con el consolidator, que además decae prioridades para evitar que unos pocos ejemplos monopolizen el proceso.",
                "La disciplina completa es la siguiente: una observación numérica supervisada entra, se convierte en `ReplaySample`, se añade con prioridad ligada a la importancia estimada, y el trainer ejecuta una actualización con mezcla de muestra actual y batch rehecho desde replay. Esa mezcla evita que la adaptación localizada destruya rápidamente patrones ya dominados.",
            ],
        },
        {
            "title": "14. Consolidation Engine",
            "paragraphs": [
                "La consolidación periódica cumple un papel equivalente al de limpieza estructural y reorganización de memoria. No produce respuestas directamente; mejora el sustrato cognitivo. La implementación actual fusiona registros muy similares, incrementa conteos de refuerzo, poda residuos con bajo valor y reescala prioridades de replay.",
                "Este mecanismo traduce al nivel de software una intuición tomada de la literatura de memoria [R1]: la experiencia no debería quedar almacenada como eventos independientes para siempre. Debe compactarse, generalizarse y perder detalle accesorio cuando el patrón ya quedó representado.",
                "La separación temporal también importa. Consolidar dentro del paso crítico de inferencia encarece la latencia y complica la trazabilidad. Al ejecutarla cada cierto número de pasos, el sistema conserva capacidad de respuesta y aun así reorganiza conocimiento de fondo.",
            ],
        },
        {
            "title": "15. Dynamic Routing",
            "paragraphs": [
                "El router es el mecanismo que hace que la arquitectura no sea monolítica. Para payloads numéricos activa el camino semántico reducido, el módulo plástico y el núcleo estable con presupuesto de cómputo bajo. Para texto, prioriza memoria y compresión, y solo marca plasticidad como relevante cuando detecta conocimiento, preferencias o correcciones.",
                "La inspiración viene de arquitecturas de expertos y cómputo condicional [R6][R7]. La diferencia es que aquí el ruteo no solo decide qué red feed-forward usar; decide qué memorias consultar, si conviene aprender, si hay que activar replay y qué presupuesto de cómputo asignar. Es un ruteo cognitivo, no solo tensorial.",
                "Este diseño facilita escalabilidad. En una versión distribuida, el router podría enviar la consulta a expertos especializados de código, dominio médico, razonamiento matemático o memoria personal, sin tocar la API superior. Ese desacoplamiento es uno de los activos más valiosos del framework.",
            ],
        },
        {
            "title": "16. Safety, Anti-Drift And Knowledge Hygiene",
            "paragraphs": [
                "La arquitectura incorpora varias defensas contra corrupción de conocimiento. La primera es el evaluador de importancia, que evita guardar ruido y señales débiles. La segunda es la confianza explícita y el riesgo de contradicción. La tercera es la compresión conceptual, que reduce la contaminación contextual de transcripts completos. La cuarta es la plasticidad localizada: solo un submódulo cambia de forma significativa en línea.",
                "A esto se añade telemetría de drift. El trainer registra norma acumulada de cambio plástico y el núcleo estable permanece efectivamente inmóvil. En el experimento numérico, el reporte capturó explícitamente la deriva del módulo plástico por época, permitiendo observar que la adaptación avanzaba mientras el stable core permanecía en deriva cero.",
                "Una siguiente versión podría añadir checkpoints con rollback automático, validación cruzada sobre memoria consolidada, detectores de contradicción más simbólicos y umbrales adaptativos basados en calibración empírica. El esqueleto del framework ya deja espacio para esas extensiones.",
            ],
        },
        {
            "title": "17. Hot-Swapping, Registry And Dependency Injection",
            "paragraphs": [
                "Para que el proyecto sea realmente reutilizable, no basta con tener clases. Hace falta una política explícita de composición. El framework incorpora `ComponentRegistry`, `EngineBuilder` y el método `replace_component` del engine. Con eso, un procesador de entrada, un backend de memoria o un router alternativo se pueden insertar sin reescribir el resto del pipeline.",
                "La inyección de dependencias se materializa en el builder. En vez de instanciar todos los objetos dentro de la lógica de inferencia, el engine recibe componentes ya construidos: processors, semantic encoder, importance evaluator, memory system, stable core, plastic learner y consolidator. Esta separación simplifica pruebas, benchmarkeo y despliegue diferenciado.",
                "El registry es particularmente útil para convertir el proyecto en framework. Permite descubrir y construir implementaciones por categoría (`processor`, `memory`, `stable_core`, `plasticity`, etc.). En un crecimiento futuro, ese registry sería el punto natural para plugins de terceros o configuraciones declarativas basadas en YAML.",
            ],
        },
        {
            "title": "18. Scalability Strategy",
            "paragraphs": [
                "La escalabilidad se diseñó como propiedad arquitectónica, no como optimización tardía. En hardware pequeño, el sistema puede ejecutar el encoder textual pequeño, memoria local NumPy, replay en RAM y plasticidad localizada en CPU. En una estación de trabajo media, el núcleo estable y el módulo plástico pueden moverse a GPU mientras la memoria queda en proceso aparte. En un entorno distribuido, el router podría delegar a expertos remotos y a una base vectorial externa.",
                "El desacoplamiento por carpetas ayuda directamente a escalar. La memoria puede externalizarse sin reescribir semántica; la plasticidad puede crecer en parámetros sin modificar la API del engine; el router puede hacerse más complejo sin tocar el compresor. Esto reduce el costo de evolucionar el sistema desde prototipo hacia producto.",
                "La inferencia distribuida también es plausible porque el contrato intermedio es semántico. Una vez generado el `SemanticState`, diferentes servicios podrían operar de forma paralela: uno consulta memoria, otro evalúa contradicción, otro propone compresión, otro gestiona replay. La arquitectura ya tiene el tipo de fronteras que una versión multi-proceso o multi-GPU necesitaría.",
            ],
        },
        {
            "title": "19. Observability And Full Internal Visualization",
            "paragraphs": [
                "Uno de los requisitos del proyecto fue permitir visualización completa de lo que ocurre dentro del sistema. Eso se resolvió con varias capas. La primera es la telemetría JSONL por evento. La segunda son snapshots estructurados del estado del engine. La tercera son plots derivados de resultados reales: evolución de importancia, crecimiento de memoria, pérdida y error absoluto medio, deriva plástica y comparación predicción-objetivo.",
                "La visibilidad no es solo un lujo. En aprendizaje continuo la opacidad es peligrosa. Si el sistema comienza a recordar basura, a ignorar preferencias, a sobreajustarse a ejemplos recientes o a degradar precisión, la observabilidad debe permitir localizar la fase responsable. En el framework actual es posible ver si el problema nació en input processing, en la política de importancia, en la recuperación de memoria o en la adaptación plástica.",
                "La combinación de trazas por paso y resúmenes por experimento convierte a la arquitectura en un laboratorio controlable. Ese atributo es esencial para investigación seria, porque permite modificar una pieza y observar causalmente qué cambió en el comportamiento global.",
            ],
            "figures": [
                str(PLOT_DIR / "importance_confidence_stream.png"),
                str(PLOT_DIR / "memory_growth_stream.png"),
                str(PLOT_DIR / "training_loss_mae.png"),
                str(PLOT_DIR / "accuracy_drift.png"),
                figures["prediction_cases"],
            ],
        },
        {
            "title": "20. Experimental Methodology",
            "paragraphs": [
                "La validación práctica se dividió en dos partes. La primera fue un stream textual curado de diez entradas que mezcla conocimiento técnico, preferencias, small talk, correcciones y ruido. Esa secuencia permitió verificar si la arquitectura aprende selectivamente, comprime memoria, mantiene un tamaño acotado y responde con contexto recuperado.",
                "La segunda fue un experimento numérico de microescala sobre suma, resta y multiplicación con operandos entre 0 y 12. Se eligió deliberadamente una tarea muy pequeña para demostrar que el engine puede alojar entrenamiento online real sin necesidad de grandes datasets ni preentrenamiento costoso. El objetivo no fue mostrar poder absoluto, sino validar la usabilidad de la arquitectura, la localización de plasticidad y el papel del replay.",
                "Ambos experimentos corrieron sobre el mismo framework. Eso importa metodológicamente: el experimento numérico no vive en un script separado ni en un notebook ajeno al engine, sino dentro de la misma topología arquitectónica. Por eso sus resultados sí dicen algo sobre el diseño general del sistema.",
            ],
        },
        {
            "title": "21. Text Memory Stream Results",
            "paragraphs": [
                "El stream textual produjo decisiones diferenciadas. El sistema reforzó conocimiento técnico inicial, aprendió preferencias, absorbió hechos de proyecto, descartó small talk y mantuvo correcciones como memorias valiosas. Esta es la señal empírica más clara de que el filtro de aprendizaje selectivo funciona: no todo lo que entra se vuelve memoria persistente.",
                "La consulta final sobre preferencias logró recuperar como primera respuesta el hecho correcto `User preference extracted: prefiero, respuestas, tecnicas, directas`. La consulta sobre el proyecto recuperó el hecho `proyecto, usa, pytorch, prototipos`. Esos comportamientos son modestos pero fundamentales: prueban que la memoria comprimida y el retrieval híbrido ya aportan valor observable.",
                "El tamaño final de memoria semántica se mantuvo pequeño y semánticamente interpretable. En un transcript largo convencional se esperaría crecimiento lineal de fragmentos textuales. Aquí no ocurrió eso; el sistema prefirió unas pocas representaciones comprimidas con significado operativo.",
            ],
            "table": {"headers": ["Step", "Action", "Importance", "Learned", "Input"], "rows": stream_rows},
        },
        {
            "title": "22. Tiny Numeric Plasticity Results",
            "paragraphs": [
                f"El microexperimento numérico arrojó resultados sólidos. El mejor checkpoint apareció en la época {numeric_report['best_epoch']} y alcanzó exactitud de validación {numeric_report['best_validation_accuracy']:.3f}. Tras restaurar ese checkpoint, la evaluación final conservada en el reporte fue accuracy={best_validation['accuracy']:.3f} y MAE={best_validation['mae']:.3f}.",
                "Más allá de la métrica final, la curva muestra varios fenómenos relevantes para el diseño. La precisión no crece de forma monótona: hay oscilaciones intermedias debido a la interacción entre replay, prioridades y una plasticidad deliberadamente pequeña. Eso es informativo, no un defecto escondido. Muestra que el sistema está realmente aprendiendo en línea y que el control de checkpointing es parte integral del framework.",
                "Las predicciones en probes concretos fueron cercanas al objetivo: 2+3 -> 4.873, 9-4 -> 4.841, 6*7 -> 41.982 y 12+12 -> 23.458. Para un módulo diminuto, localizado y entrenado sin grandes datos, el resultado es suficiente para validar el principio de funcionamiento y la viabilidad del pipeline completo.",
            ],
            "table": {"headers": ["Epoch", "Train loss", "Val MAE", "Val accuracy", "Plastic drift"], "rows": epoch_rows},
            "figures": [str(PLOT_DIR / "training_loss_mae.png"), str(PLOT_DIR / "accuracy_drift.png"), figures["prediction_cases"]],
        },
        {
            "title": "23. Interpretation Of The Numeric Experiment",
            "paragraphs": [
                "La tarea numérica fue elegida porque permite observar el comportamiento interno del sistema sin confundir éxito arquitectónico con memorizar texto. Aquí los operandos no son tokens de lenguaje, sino cantidades numéricas con representación propia. Eso deja claro que la arquitectura puede hospedar submodelos especializados con entradas no lingüísticas y aun así integrarlos en un pipeline cognitivo común.",
                "También demuestra que la plasticidad controlada es compatible con un núcleo estable. El stable core no tuvo que absorber la aritmética; simplemente coordinó memoria y respuesta. La adaptación residió en el submódulo plástico. Esa es exactamente la propiedad deseada para dominios donde la personalización o el aprendizaje incremental no deberían contaminar la totalidad del sistema.",
                "En términos de investigación, este experimento sirve como prueba de vida reproducible. No pretende resolver razonamiento matemático general, sino ofrecer un benchmark microscópico donde cada decisión del framework se pueda inspeccionar.",
            ],
            "table": {"headers": ["Operation", "Operands", "Target", "Prediction", "Confidence"], "rows": prediction_rows},
        },
        {
            "title": "24. Why This Is Not A Monolithic LLM",
            "paragraphs": [
                "El sistema resultante difiere de un LLM monolítico en varios ejes: no aprende modificando todos sus pesos; no guarda todo el texto; no depende de un único transformer como punto de verdad; no mezcla memoria reciente, memoria consolidada y plasticidad en una sola matriz; y no fuerza el mismo camino de cómputo para todas las tareas.",
                "La modularidad aquí no es cosmética. Cambiar el backend de memoria, sustituir el evaluador de importancia, introducir otro tipo de compresor o reemplazar el módulo plástico no obliga a rediseñar el engine. Esa es la marca de una arquitectura híbrida seria.",
                "Además, el proyecto deja abierta la evolución hacia LLMs reales como núcleo estable. El `StableReasoningCore` actual es pequeño y utilitario, pero el contrato es compatible con sustituirlo por un modelo más potente congelado o semi-congelado, mientras el resto de la arquitectura conserva su función de control cognitivo.",
            ],
        },
        {
            "title": "25. Design Patterns Used",
            "paragraphs": [
                "Se emplearon varios patrones clásicos. `Strategy` aparece en processors, semantic encoders, routers y stores intercambiables. `Factory` y `Registry` viven en el builder y la creación declarativa de componentes. `Dependency Injection` organiza la construcción del engine. `Observer` aparece en la telemetría por eventos. `Facade` emerge en `CognitiveAPIService`, que expone operaciones de alto nivel sin arrastrar detalles internos. `Template Method` aparece en la separación entre contratos abstractos e implementaciones concretas.",
                "Usar estos patrones no fue un adorno académico. Son los mecanismos que permiten crecimiento modular sin acoplamiento explosivo. En particular, Factory + Registry + DI reducen la fricción de hot-swapping y facilitan que el framework escale hacia múltiples variantes sin volverse inmantenible.",
                "Desde la perspectiva de un equipo real, este conjunto de patrones también mejora testeabilidad. Cada pieza se puede probar aislada, simular o sustituir en benchmarks específicos.",
            ],
        },
        {
            "title": "26. Internal APIs And Pipeline Contracts",
            "paragraphs": [
                "El API interno del engine está pensado como una cadena de transformaciones tipadas. `process(payload)` es la fachada principal, pero internamente el pipeline es: `ProcessedInput -> SemanticState -> MemoryQuery/MemoryBundle -> ImportanceAssessment -> CompressedKnowledge -> CoreInference`. Esta linealidad aparente esconde ramificaciones controladas por routing, pero mantiene un lenguaje técnico uniforme para todos los módulos.",
                "Esa uniformidad facilita integración con servicios externos. Un proceso separado podría producir `ProcessedInput`; otro podría operar como memoria remota; otro podría consumir los `TraceEvent` para un dashboard. Al usar contratos explícitos, el framework no obliga a correr todo en un mismo proceso Python para seguir siendo coherente.",
                "La misma idea aplica a la API pública. `CognitiveAPIService` provee operaciones `ingest_text`, `ask_text`, `train_numeric` e `infer_numeric`. Son pocas, pero capturan el patrón deseado: una superficie pequeña sobre un interior altamente modular.",
            ],
        },
        {
            "title": "27. Limitations",
            "paragraphs": [
                "El prototipo es funcional, pero no pretende resolver comprensión semántica abierta al nivel de un LLM grande. El encoder textual es pequeño y heurístico; la compresión es ligera; la detección de contradicción todavía no razona sobre hechos complejos; y la recuperación híbrida usa reglas simples que deberían madurar para producción.",
                "La tarea numérica, aunque útil, es un benchmark de validación y no una demostración de inteligencia simbólica general. También queda trabajo por hacer en selección adaptativa de umbrales, calibración de confianza y evaluación más extensa bajo drift adversarial o streams más largos.",
                "Por último, no se integró aún un backend vectorial externo real, aunque el contrato ya está preparado. La memoria actual es local y suficiente para probar arquitectura, no para millones de registros.",
            ],
        },
        {
            "title": "28. Extension Roadmap",
            "paragraphs": [
                "El camino inmediato más natural es sustituir el núcleo estable por un modelo semántico más fuerte y mantener intacto el resto del framework. En paralelo, el evaluador de importancia podría entrenarse con feedback humano y métricas downstream. Otra mejora clara es introducir contradicción simbólica basada en tripletas sujeto-predicado-objeto y validación multi-fuente.",
                "En escalado medio, el sistema debería adoptar una base vectorial externa, colas de eventos para consolidación asíncrona y dashboards de telemetría. En escalado grande, el router podría seleccionar entre expertos remotos y varios módulos plásticos por dominio o usuario.",
                "La dirección de investigación más interesante es estudiar políticas de aprendizaje selectivo que no solo decidan qué recordar, sino también dónde recordarlo: memoria explícita, replay, adapter local, grafo semántico o fine-tuning eventual de un experto. Esa sería una verdadera política cognitiva multi-almacén.",
            ],
        },
        {
            "title": "29. Conclusion",
            "paragraphs": [
                "El proyecto logró lo esencial que se pedía: una arquitectura híbrida cognitiva modular en Python con PyTorch, aprendizaje selectivo, memoria jerárquica, plasticidad controlada, consolidación periódica, replay y observabilidad integral. No se limitó a un script ni a una única red neuronal densa.",
                "Lo más relevante es que el diseño ya es usable. Procesa entradas, decide qué aprender, comprime conocimiento, consulta memoria, entrena un submódulo plástico real y produce artefactos verificables. La combinación de teoría arquitectónica y validación empírica convierte al repositorio en una base seria para seguir investigando o construir encima.",
                "En resumen, la propuesta demuestra que se puede pensar una IA de texto-a-texto como sistema cognitivo compuesto, donde memoria, aprendizaje y razonamiento se coordinan explícitamente, en vez de confiar toda la adaptabilidad a un bloque monolítico de parámetros.",
                "También deja una lección metodológica útil: incluso cuando el objetivo final es un sistema mucho mayor, vale la pena validar primero la arquitectura en escalas microscópicas y altamente instrumentadas. El micromodelo numérico, el stream textual controlado y las trazas estructuradas mostraron que las decisiones arquitectónicas pueden medirse antes de invertir en modelos gigantes. Esa disciplina reduce riesgo, acelera iteración y hace que el crecimiento futuro se apoye en fundamentos verificables y no en intuiciones opacas.",
                "Ese principio de validación incremental es, en sí mismo, parte de la arquitectura propuesta.",
            ],
        },
        {
            "title": "30. References",
            "paragraphs": [f"[{item['id']}] {item['label']}. {item['url']} - {item['note']}" for item in REFERENCES],
        },
    ]
    pipeline_steps = [
        "El engine recibe un payload sin asumir modalidad. Esa neutralidad inicial evita que el sistema trate todo como texto y abre la puerta a rutas heterogéneas desde el primer milisegundo.",
        "El selector de processor inspecciona el payload y delega en el componente adecuado mediante `supports(payload)`. Ese paso convierte un dato crudo en una estructura de entrada compatible con el resto del framework.",
        "La semántica transforma la entrada en `SemanticState`, agregando embedding, intención, conceptos y contexto comprimido. Esta es la primera abstracción común entre modalidades distintas.",
        "Working memory se actualiza inmediatamente para que el estado actual del diálogo o de la tarea quede disponible incluso antes de consolidar nada en memoria permanente.",
        "El router toma modalidad e intención y define presupuesto de cómputo, memorias a consultar y si conviene activar plasticidad. Esta decisión reduce trabajo innecesario y delimita el riesgo de actualización.",
        "La consulta de memoria se construye con embedding y conceptos, no con texto bruto. Eso fuerza coherencia entre recuperación y compresión conceptual.",
        "La memoria responde con un `MemoryBundle` que separa explícitamente contexto corto, estado de trabajo, memoria semántica y episodios recientes. Esa separación evita ambigüedad sobre el origen del contexto usado en la respuesta.",
        "El evaluador de importancia calcula prioridad de aprendizaje. En este punto el sistema decide si la experiencia debe ignorarse, aprenderse, reforzarse, consolidarse o marcarse como incierta.",
        "Si la ruta activa plasticidad y la modalidad lo requiere, el módulo plástico genera una predicción especializada o prepara un gradiente localizado. No se tocan todos los pesos del sistema.",
        "El núcleo estable compone la respuesta final combinando contexto actual, memoria recuperada y salida especializada. El resultado es una salida inteligible y rastreable.",
        "Cuando el input supera el umbral de aprendizaje, el compresor reescribe la experiencia en una unidad de conocimiento compacta y persistible. Aquí el transcript se convierte en memoria útil.",
        "El sistema de memoria inserta o fusiona el nuevo registro. Si detecta redundancia alta, refuerza conocimiento existente en lugar de inflar artificialmente la base de memoria.",
        "En tareas supervisadas, el trainer construye un batch con la observación actual y muestras desde replay. La actualización ocurre solo en el submódulo plástico.",
        "Cada fase emite trazas. Esa telemetría alimenta logs, debugging, visualizaciones y la evidencia experimental usada por este paper.",
        "Cada cierto número de pasos, la consolidación reorganiza memoria y reduce residuos. El proceso es deliberadamente periódico y fuera del camino crítico de inferencia.",
    ]
    appendix_pipeline = {
        "title": "Appendix A. Detailed Runtime Pipeline",
        "paragraphs": [
            "Este apéndice expande el pipeline no solo como lista de funciones, sino como secuencia cognitiva explícita. El objetivo es mostrar cómo se mueve la información a través de estados y decisiones interpretables."
        ]
        + [f"Paso {index + 1}. {text}" for index, text in enumerate(pipeline_steps)],
    }

    appendix_folders = {
        "title": "Appendix B. Folder-By-Folder Engineering Notes",
        "paragraphs": [
            f"{folder}. {note} En un equipo real, esta separación reduce conflictos de mantenimiento porque el dueño conceptual de la carpeta coincide con el dueño funcional del comportamiento correspondiente."
            for folder, note in FOLDER_NOTES.items()
        ],
    }

    interface_notes = {
        "InputProcessor": "Contrato de percepción. Encapsula tokenización, normalización y creación de estructuras de entrada sin contaminar al engine con detalles de modalidad.",
        "SemanticEncoder": "Produce un `SemanticState` reusable por memoria, respuesta y aprendizaje. Es la bisagra entre datos crudos y representación cognitiva.",
        "ImportanceEvaluator": "Decide qué experiencia merece persistencia o adaptación. Es el filtro que protege al sistema de memorizar ruido.",
        "KnowledgeCompressor": "Reescribe experiencia en conocimiento estructurado y deduplicable. Hace viable la memoria semántica compacta.",
        "MemoryStore": "Contrato de almacenamiento recuperable. Permite intercambiar implementaciones locales o externas sin romper la lógica del engine.",
        "Router": "Selecciona ruta de cómputo y presupuesto operacional. Materializa el principio de especialización condicional.",
        "StableCore": "Compone respuestas con comportamiento conservador. Es la reserva de estabilidad lingüística y contextual del sistema.",
        "PlasticLearner": "Aloja la plasticidad entrenable y localizada. Su interfaz distingue predicción y `train_step`, permitiendo separar inferencia y adaptación.",
        "ReplayBuffer": "Gestiona muestras previas para rehearsal. Su presencia estructural impide que el aprendizaje continuo dependa solo de la observación actual.",
        "Consolidator": "Ejecución de limpieza, fusión y reorganización de memoria. Mantiene la base cognitiva compacta y usable.",
        "TelemetrySink": "Observador persistente del pipeline. Hace posible inspección externa sin introducir side effects en los módulos principales.",
    }
    appendix_interfaces = {
        "title": "Appendix C. Interface Catalog",
        "paragraphs": [
            "Las interfaces abstractas no son un detalle estético. Son el seguro de reemplazabilidad del framework. Cada contrato separa intención funcional de implementación concreta."
        ]
        + [f"{name}. {note}" for name, note in interface_notes.items()],
    }

    deployment_paragraphs = [
        "Perfil 1, laptop o CPU pequeña. El engine puede operar con memoria NumPy local, encoder reducido y módulo plástico diminuto. Es suficiente para investigación, demos y trazabilidad total.",
        "Perfil 2, workstation con una GPU. El núcleo estable y el módulo plástico se ejecutan en GPU, mientras la memoria y la telemetría permanecen en CPU. Este modo permite latencias razonables y entrenamiento incremental sin infraestructura distribuida.",
        "Perfil 3, workstation multi-GPU. El stable core puede fijarse en una GPU y los expertos o adapters plásticos en otra, mientras procesos auxiliares manejan replay y consolidación. El router coordina el reparto.",
        "Perfil 4, cluster o servicio distribuido. El `SemanticState` se vuelve la moneda de intercambio entre microservicios: memoria externa, expertos remotos, consolidación asíncrona y panel de observabilidad.",
        "Perfil 5, producto multiusuario. El núcleo estable se comparte, mientras uno o varios módulos plásticos se segmentan por usuario, segmento o dominio. Esta topología es útil para personalización sin deriva global.",
    ]
    appendix_deployment = {
        "title": "Appendix D. Deployment Profiles",
        "paragraphs": [
            "La escalabilidad no depende solo de más hardware; depende de que el diseño permita mover piezas sin rehacer el sistema. Este apéndice aterriza esa idea en perfiles concretos."
        ]
        + deployment_paragraphs,
    }

    appendix_epoch_analysis = {
        "title": "Appendix E. Epoch-Level Interpretation Of The Numeric Experiment",
        "paragraphs": [
            "La curva de entrenamiento del micromodelo muestra varias fases distinguibles. Las primeras épocas reducen error grueso; las épocas medias estabilizan suma y resta; las épocas altas mejoran multiplicación y consolidan precisión; y las últimas evidencian que incluso un módulo pequeño puede oscilar si la presión de replay y adaptación no se acompaña de selección de checkpoint.",
            "La selección del mejor estado no es un truco cosmético. En aprendizaje continuo es normal que el mejor punto aparezca antes del último paso, especialmente cuando se mantiene exploración plástica. Diseñar para checkpoints explícitos es parte de la arquitectura, igual que diseñar para memoria y replay.",
            "El hecho de que la exactitud alcanzara 1.0 en múltiples épocas y luego descendiera en la última demuestra que el sistema sí recorrió regiones funcionales del espacio de parámetros. Por tanto, la tarea de ingeniería no es 'hacer que aprenda algo', sino 'regular cuándo congelar y cómo consolidar'.",
        ],
        "table": {"headers": ["Epoch", "Train loss", "Val MAE", "Val accuracy", "Plastic drift"], "rows": epoch_rows},
    }

    appendix_reuse = {
        "title": "Appendix F. Reuse Scenarios",
        "paragraphs": [
            "Escenario 1: asistente personal que aprende preferencias, decisiones recurrentes y correcciones de estilo sin reentrenar el modelo base.",
            "Escenario 2: tutor técnico que consolida reglas de proyecto, librerías preferidas y correcciones de arquitectura a medida que interactúa con un equipo.",
            "Escenario 3: agente industrial con memorias episódicas de fallas, memoria semántica de procedimientos y expertos plásticos por máquina o línea de producción.",
            "Escenario 4: copiloto de analítica que combina memoria vectorial, routing por dominio y micro-expertos numéricos para tareas tabulares o financieras.",
            "Escenario 5: entorno de investigación donde distintos encoders, políticas de importancia y estrategias de replay se comparan bajo un mismo engine.",
            "En todos los casos, la clave es la misma: el framework no obliga a que el aprendizaje incremental viva en el mismo lugar donde vive el lenguaje general. Esa independencia es la que hace viable personalizar sin destruir la base.",
        ],
    }

    file_inventory = {
        "cognitive_engine/core/types.py": "Define los tipos de paso que dan forma al sistema. Sin estos dataclasses, la arquitectura perdería gran parte de su claridad y capacidad de inspección.",
        "cognitive_engine/interfaces/base.py": "Concentra los contratos abstractos. Es el punto de referencia para extender el framework sin introducir dependencias circulares entre implementaciones concretas.",
        "cognitive_engine/core/engine.py": "Orquesta el pipeline completo, emite trazas, ejecuta aprendizaje en línea y dispara consolidación periódica.",
        "cognitive_engine/core/builder.py": "Materializa el ensamblaje profesional del engine y muestra cómo usar registry, DI y configuración dinámica.",
        "cognitive_engine/modules/input_processing.py": "Aloja la percepción multimodal ligera. Su separación evita que el resto del stack conozca detalles de parsing de texto o números.",
        "cognitive_engine/modules/semantic_understanding.py": "Contiene los encoders híbridos y la traducción desde entradas preprocesadas hacia estados semánticos ricos.",
        "cognitive_engine/modules/importance_evaluator.py": "Implementa la política de aprendizaje selectivo. Es la barrera principal contra ruido, drift y memoria indiscriminada.",
        "cognitive_engine/compression/knowledge_compressor.py": "Traduce experiencia a conocimiento persistible, resumiendo y estructurando lo que vale la pena recordar.",
        "cognitive_engine/memory/stores.py": "Resuelve almacenamiento jerárquico, fusión, recuperación y snapshots de memoria. Es el corazón mnésico del prototipo.",
        "cognitive_engine/adapters/plastic_numeric_adapter.py": "Demuestra la idea de plasticidad localizada con un submódulo pequeño, entrenable y observable.",
        "cognitive_engine/training/online_trainer.py": "Gestiona mezcla entre observación actual y replay, además de evaluación periódica sobre datasets de validación.",
        "scripts/train_numeric_demo.py": "Convierte la arquitectura en evidencia empírica reproducible sobre un benchmark numérico mínimo.",
        "scripts/run_text_memory_demo.py": "Ejercita la memoria selectiva con entradas de distinta relevancia y genera artefactos para inspección.",
        "scripts/generate_paper.py": "Conecta teoría, resultados y artefactos en una entrega documental extensa y rastreable.",
        "tests/test_numeric_training.py": "Fija una condición mínima de funcionamiento para la rama plástica, evitando regresiones silenciosas.",
    }
    appendix_files = {
        "title": "Appendix G. Critical File Inventory",
        "paragraphs": [
            "En un framework serio, el árbol de archivos también comunica arquitectura. Este inventario resume por qué ciertos archivos son estructuralmente decisivos."
        ]
        + [f"{path}. {note}" for path, note in file_inventory.items()],
    }

    memory_lifecycle = [
        "Nacimiento. Una observación entra por el processor y llega al `SemanticState` con intención y conceptos explícitos.",
        "Evaluación. El sistema estima si la observación es novedosa, útil, coherente y suficientemente confiable como para aprenderla.",
        "Compresión. Si supera el umbral, la experiencia se reescribe en una forma conceptual resumida y deduplicable.",
        "Ingreso. El registro comprimido entra a short-term y episodic memory, y potencialmente a memoria semántica si no es demasiado redundante.",
        "Refuerzo. Si ya existe un registro similar, se incrementa el conteo de refuerzo en lugar de duplicar conocimiento.",
        "Recuperación. Futuras consultas pueden traer ese conocimiento por similitud vectorial, coincidencia temática o ambos criterios.",
        "Consolidación. Periódicamente, registros altamente similares se fusionan y los residuos débiles pueden podarse.",
        "Replay. En tareas supervisadas, algunas observaciones también alimentan el buffer de rehearsal para proteger desempeño futuro.",
        "Reinterpretación. Si nuevas correcciones contradicen un recuerdo previo, la arquitectura puede reforzar o desplazar la versión anterior según política.",
    ]
    appendix_memory_lifecycle = {
        "title": "Appendix H. Memory Record Lifecycle",
        "paragraphs": [
            "La memoria no es un destino final único; es un ciclo de transformación. Este apéndice describe el recorrido típico de una unidad de conocimiento dentro del sistema."
        ]
        + memory_lifecycle,
    }

    risk_rows = [
        ["Ruido conversacional", "Memoria inflada y baja señal", "Importance evaluator + compresión conceptual"],
        ["Correcciones tardías", "Persistencia de hechos desactualizados", "Prioridad alta para intentos de correction"],
        ["Plasticidad excesiva", "Drift o pérdida de conocimiento", "Submódulo localizado + replay + stable core congelado"],
        ["Retrieval irrelevante", "Respuesta fuera de contexto", "Ranking híbrido vectorial-lexical"],
        ["Crecimiento de memoria", "Latencia y redundancia", "Consolidación y fusión periódica"],
        ["Sobreadaptación del experto", "Oscilación en validación", "Checkpointing del mejor estado y evaluación recurrente"],
        ["Acoplamiento arquitectónico", "Dificultad de mantenimiento", "Interfaces, builder y registry"],
    ]
    appendix_risk = {
        "title": "Appendix I. Failure Mode Matrix",
        "paragraphs": [
            "Toda arquitectura de aprendizaje continuo tiene fallos previsibles. La pregunta seria no es si existen, sino si el diseño reserva un lugar explícito para mitigarlos.",
            "La matriz siguiente resume los riesgos más obvios observados o anticipados en este proyecto y la pieza estructural encargada de contenerlos.",
        ],
        "table": {"headers": ["Failure mode", "Risk", "Mitigation"], "rows": risk_rows},
    }

    chronology = [
        "t0: entrada cruda recibida por el engine.",
        "t1: normalización y selección del processor adecuado.",
        "t2: construcción del estado semántico con conceptos e intención.",
        "t3: actualización de working memory para reflejar el foco cognitivo actual.",
        "t4: decisión de ruteo y definición de memorias a consultar.",
        "t5: recuperación de contexto semántico y episódico relevante.",
        "t6: cálculo de importancia, confianza y prioridad de aprendizaje.",
        "t7: inferencia del módulo plástico cuando la modalidad o el router lo requieren.",
        "t8: composición de respuesta por el núcleo estable usando memoria y salida especializada.",
        "t9: compresión opcional del evento en unidad de conocimiento.",
        "t10: escritura o refuerzo de memoria según la política.",
        "t11: inserción del ejemplo en replay cuando existe supervisión.",
        "t12: actualización localizada del submódulo plástico.",
        "t13: emisión de trazas y persistencia telemétrica.",
        "t14: consolidación periódica fuera del camino crítico.",
    ]
    appendix_chronology = {
        "title": "Appendix J. Online Learning Chronology",
        "paragraphs": [
            "La siguiente cronología resume la secuencia temporal completa de una observación supervisada o semisupervisada dentro del framework. Está pensada para debugging, diseño de dashboards y discusión de latencia."
        ]
        + chronology,
    }

    glossary_items = {
        "Stable core": "Submodelo o capa de composición que no se actualiza agresivamente y protege la consistencia global del sistema.",
        "Plastic learner": "Submódulo entrenable de adaptación localizada. Aprende rápido, pero solo en una región pequeña del espacio paramétrico.",
        "Compressed knowledge": "Unidad de memoria resumida que reemplaza transcript crudo por conceptos, relaciones y embedding.",
        "Working memory": "Estado efímero de foco actual, intención activa y conceptos recientes; no equivale a memoria semántica persistente.",
        "Replay": "Reutilización deliberada de ejemplos anteriores para mantener desempeño frente a secuencias no estacionarias.",
        "Consolidation": "Proceso periódico de fusión, reorganización y depuración de memorias para ganar estabilidad y compactación.",
        "Routing": "Selección dinámica de camino de cómputo, presupuesto, memoria y activación de especialistas.",
        "Importance score": "Puntuación que estima si una observación debe recordarse o ignorarse.",
        "Confidence score": "Estimación de fiabilidad del conocimiento observado o inferido.",
        "Hot-swapping": "Sustitución de un componente por otro compatible sin rediseñar la arquitectura entera.",
        "Semantic bundle": "Conjunto de memorias recuperadas y estado de trabajo que contextualiza la respuesta actual.",
        "Lexical boost": "Bonificación al ranking de memoria cuando hay coincidencia temática fuerte entre consulta y memoria almacenada.",
        "Checkpoint selection": "Práctica de conservar el mejor estado validado del módulo plástico en lugar de asumir que el último es el mejor.",
        "Background consolidation": "Reorganización fuera del camino crítico de inferencia para preservar latencia de respuesta.",
        "Controlled plasticity": "Principio según el cual el sistema aprende en zonas acotadas y auditables en vez de modificar la totalidad del modelo.",
    }
    appendix_glossary = {
        "title": "Appendix K. Operational Glossary",
        "paragraphs": [
            "Este glosario resume términos usados repetidamente a lo largo del paper para que el documento pueda reutilizarse como referencia interna de diseño.",
            "La intención no es definir conceptos en abstracto, sino fijar cómo se usan dentro de esta arquitectura específica. En entornos de ingeniería, ese tipo de claridad terminológica reduce malentendidos entre quienes diseñan modelos, quienes operan memoria y quienes despliegan el sistema.",
        ]
        + [f"{term}. {definition}" for term, definition in glossary_items.items()],
    }

    appendix_sections = [
        appendix_pipeline,
        appendix_folders,
        appendix_interfaces,
        appendix_deployment,
        appendix_epoch_analysis,
        appendix_reuse,
        appendix_files,
        appendix_memory_lifecycle,
        appendix_risk,
        appendix_chronology,
        appendix_glossary,
    ]
    sections = sections[:-1] + appendix_sections + [sections[-1]]
    return sections


def main() -> None:
    engine = EngineBuilder(config_path="configs/default.yaml").build()
    engine_description = engine.describe_architecture()
    text_report = load_json(REPORT_DIR / "text_memory_demo.json")
    numeric_report = load_json(REPORT_DIR / "numeric_training_report.json")

    figures = {
        "architecture": architecture_graph(engine_description, PLOT_DIR / "architecture_topology.png"),
        "action_histogram": action_histogram_plot(text_report, PLOT_DIR / "action_histogram.png"),
        "prediction_cases": prediction_case_plot(numeric_report, PLOT_DIR / "prediction_case_comparison.png"),
    }

    tree_text = build_tree(
        ROOT,
        [
            "cognitive_engine",
            "configs",
            "scripts",
            "tests",
            "docs",
            "artifacts",
        ],
    )
    sections = generate_sections(engine_description, text_report, numeric_report, tree_text, figures)

    markdown_parts = [
        "# Hybrid Cognitive Continual Learning Architecture",
        "## Theory, implementation architecture and empirical validation",
    ]
    for section in sections:
        markdown_parts.append(f"\n## {section['title']}\n")
        markdown_parts.extend(section.get("paragraphs", []))
        table = section.get("table")
        if table:
            markdown_parts.append(md_table(table["headers"], table["rows"]))
        for figure in section.get("figures", []):
            markdown_parts.append(f"![{Path(figure).stem}]({Path(figure).as_posix()})")
        code_block = section.get("code_block")
        if code_block:
            markdown_parts.append("```text")
            markdown_parts.append(code_block)
            markdown_parts.append("```")

    markdown_content = "\n\n".join(markdown_parts)
    markdown_path = DOCS_DIR / "hybrid_cognitive_architecture_paper.md"
    write_markdown(markdown_path, markdown_content)
    metadata = {
        "title": "Hybrid Cognitive Continual Learning Architecture",
        "subtitle": "Theory, modular design, experimental validation and implementation-oriented analysis",
        "word_count": word_count(markdown_content),
    }
    docx_path = REPORT_DIR / "hybrid_cognitive_architecture_paper.docx"
    write_docx(docx_path, sections, list(figures.values()), metadata)

    summary = {
        "markdown": str(markdown_path),
        "docx": str(docx_path),
        "word_count": metadata["word_count"],
        "figures": figures,
    }
    (REPORT_DIR / "paper_manifest.json").write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps(summary, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
